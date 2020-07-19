import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from form import FileForm

"""

This python file is better implementaion of the File Template Creator script as here the Document object and the related methods are
Encapsulated, so the document object doesn't have to passed from one method to another

"""

class PracFile:
    
    def __init__(self): 
        self.document = Document()

    def create_index_table(self, num_prac):    

        # This method Creates the index table   

        para = self.document.add_paragraph()

        para.alignment= WD_ALIGN_PARAGRAPH.CENTER #The Word "Index"
        run = para.add_run("INDEX\n")
        font = run.font
        font.name = "Arial"
        font.bold = True
        font.size = Pt(20)

        table = self.document.add_table(rows=num_prac+1, cols=3)
        table.cell(0,0).text="S.No"
        table.cell(0,1).text="Program"
        table.cell(0,2).text="Date"

        for i in range(1,num_prac+1):
            table.cell(i,0).text=f"{i}"

        self.document.add_page_break() # To to move to the next page
        


    def create_exp(self, num_prac, var_list, file_name):

        self.create_index_table(num_prac)

        for i in range(1,num_prac+1): 
            for part in var_list: # This inner loop is iterated for all the parts of an experiment(aim,code etc)

                para = self.document.add_paragraph() 

                # Checking the Alignment
                if(part.alignment.data == "Center"): 
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    para.alignment= WD_ALIGN_PARAGRAPH.LEFT
                
                # For the aim the experiment number also needs to be mentioned                
                if(part.message == "Enter the aim of the experiment\n"):
                    run = para.add_run(f"{i}. { part.message }")
                else:
                    run = para.add_run(f"{ part.message }")

                # Applying various font properties
                font = run.font
                font.name = part.font.data
                font.bold = part.bold.data
                font.italic = part.italics.data
                font.underline = part.underline.data
                font.size = Pt(int(part.size.data))

            self.document.add_page_break()
        self.document.save(os.path.dirname(__file__) + f"/tmp/uploads/{file_name}")

        # Here os.path.dirname is used so that we can reach the directory where this python script is present
